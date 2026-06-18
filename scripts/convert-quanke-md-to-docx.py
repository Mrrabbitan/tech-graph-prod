#!/usr/bin/env python3
"""Convert output/quanke-agent-decision-plan-exec.md to .docx with embedded diagram PNG."""
from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.shared import Cm, Pt

ROOT = os.path.normpath(os.path.join(os.path.dirname(__file__), ".."))
MD_PATH = os.path.join(ROOT, "output", "quanke-agent-decision-plan-exec.md")
DOCX_PATH = os.path.join(ROOT, "output", "quanke-agent-decision-plan-exec.docx")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = row[j] if j < len(row) else ""
            table.rows[i].cells[j].text = cell


def main() -> int:
    if not os.path.isfile(MD_PATH):
        print("Missing", MD_PATH, file=sys.stderr)
        return 1
    text = open(MD_PATH, encoding="utf-8").read()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(11)

    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Menlo"
        run.font.size = Pt(9)
        code_buf = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        img_m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img_m:
            alt, rel = img_m.groups()
            path = rel if os.path.isabs(rel) else os.path.normpath(os.path.join(ROOT, rel))
            if os.path.isfile(path):
                doc.add_paragraph(alt or "图")
                doc.add_picture(path, width=Cm(16.5))
            else:
                doc.add_paragraph(f"[图片缺失: {path}]")
            i += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if set("".join(row)) <= {"-", ":"}:
                    i += 1
                    continue
                rows.append(row)
                i += 1
            add_table(doc, rows)
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.strip() == "":
            pass
        elif line.strip() == "\\newpage":
            doc.add_page_break()
        else:
            doc.add_paragraph(line)

        i += 1

    flush_code()

    doc.save(DOCX_PATH)
    print("Wrote", DOCX_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
